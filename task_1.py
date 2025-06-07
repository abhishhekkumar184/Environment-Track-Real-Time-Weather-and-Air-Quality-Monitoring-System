from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()
doc.add_heading('Final Project Report', 0)

# Title and Objective
doc.add_heading('Project Title:', level=1)
doc.add_paragraph('EnviroTrack: Real-Time Weather and Air Quality Monitoring System')

doc.add_heading('Objective:', level=1)
doc.add_paragraph(
    "To develop a real-time dashboard that fetches and visualizes weather and air quality data using public APIs. "
    "The goal is to provide users with insightful data visualization and environmental awareness for any location."
)

# Tools and Technologies
doc.add_heading('Tools and Technologies Used:', level=1)
doc.add_paragraph('Frontend: HTML, CSS, JavaScript (or React.js)')
doc.add_paragraph('Backend: Python (Flask) or Node.js (Express)')
doc.add_paragraph('Data Visualization: Chart.js / Plotly / Streamlit')
doc.add_paragraph('APIs: OpenWeatherMap API, IQAir API')
doc.add_paragraph('Deployment: Streamlit Cloud / Render / GitHub Pages')

# Features
doc.add_heading('Key Features:', level=1)
doc.add_paragraph('🔍 Search by City/Location')
doc.add_paragraph('🌡️ Display Real-Time Weather Info: Temperature, humidity, wind, cloud coverage')
doc.add_paragraph('🌫️ Display Air Quality (AQI) with health recommendation tags')
doc.add_paragraph('📈 Visualize Historical Trends using line/bar charts')
doc.add_paragraph('📦 Save Search History (optional)')
doc.add_paragraph('🎨 Responsive & Clean UI using Bootstrap or Tailwind CSS')

# System Design
doc.add_heading('System Design:', level=1)
doc.add_paragraph('1. User inputs city name.')
doc.add_paragraph('2. Backend fetches weather and air quality data from respective APIs.')
doc.add_paragraph('3. Parsed data is sent to frontend.')
doc.add_paragraph('4. Frontend displays real-time values and visualizes trends with charts.')

# Sample API Integration Code
doc.add_heading('Sample Python API Integration (Flask):', level=1)
doc.add_paragraph("""
import requests
from flask import Flask, request, jsonify, render_template

app = Flask(__name__)
WEATHER_API = "YOUR_OPENWEATHERMAP_KEY"
AIR_API = "YOUR_IQAIR_API_KEY"

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/get_data', methods=['POST'])
def get_data():
    city = request.form['city']
    weather_url = f"http://api.openweathermap.org/data/2.5/weather?q={city}&appid={WEATHER_API}&units=metric"
    air_url = f"http://api.airvisual.com/v2/city?city={city}&key={AIR_API}"

    weather_res = requests.get(weather_url).json()
    air_res = requests.get(air_url).json()

    return jsonify({
        'temperature': weather_res['main']['temp'],
        'humidity': weather_res['main']['humidity'],
        'wind': weather_res['wind']['speed'],
        'aqi': air_res['data']['current']['pollution']['aqius']
    })

if __name__ == '__main__':
    app.run(debug=True)
""")

# Output and Results
doc.add_heading('Expected Output:', level=1)
doc.add_paragraph('A real-time dashboard that shows:')
doc.add_paragraph('- Current temperature, humidity, and wind speed')
doc.add_paragraph('- Current AQI with health warnings')
doc.add_paragraph('- Interactive charts for AQI and temperature trends')

# Conclusion
doc.add_heading('Conclusion:', level=1)
doc.add_paragraph(
    "This project demonstrates how public APIs can be integrated into real-time applications for meaningful insights. "
    "It highlights the power of data visualization and web development to build informative dashboards."
)

# Save the document
output_path = r"C:\Users\ADMIN\Desktop\EnviroTrack_Project_Report.docx"
doc.save(output_path)
